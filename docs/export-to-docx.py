# -*- coding: utf-8 -*-
"""
Xuat docs/NGHIEP_VU_BACKEND.md (nguon su that, nam trong repo BE) ra file Word.

Markdown la NGUON, .docx la BAN XUAT. Sua tai lieu thi sua file .md roi chay lai
script nay -- khong sua thang vao .docx, neu khong hai ban se troi lech nhau.

Chay:  python docs/export-to-docx.py
Can:   pip install python-docx
Luu y: dong file .docx trong Word truoc khi chay, neu khong se bao Permission denied.

Duong dan tinh theo vi tri script nen chay tu thu muc nao cung duoc.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

DOCS_DIR = Path(__file__).resolve().parent                # ParkingManagement_BE/docs/
BE_ROOT = DOCS_DIR.parent                                 # ParkingManagement_BE/
SRC = str(DOCS_DIR / "NGHIEP_VU_BACKEND.md")
OUT = str(BE_ROOT.parent / "docs" / "PBMS_Tai_Lieu_Nghiep_Vu_Backend.docx")

CODE_FONT = "Consolas"
BODY_FONT = "Segoe UI"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
CODE_COLOR = RGBColor(0xA3, 0x15, 0x15)


def shade(element, fill):
    """To nen cho o bang / doan van."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def add_toc_field(paragraph):
    """Chen truong TOC that -- Word tu sinh muc luc khi bam Update Field."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Bam chuot phai vao day -> Update Field de sinh muc luc."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run._r.append(el)


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def add_rich_text(paragraph, text):
    """Doc **dam**, `ma nguon`, *nghieng* thanh cac run rieng."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.font.name = CODE_FONT
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_COLOR
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            paragraph.add_run(piece[1:-1]).italic = True
        else:
            paragraph.add_run(piece)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(lines))
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    shade(p._p.get_or_add_pPr(), "F4F4F4")
    return p


def is_table_sep(line):
    return bool(re.match(r"^\|[\s:\-|]+\|$", line.strip()))


BLOCK_START = re.compile(r"^(#{1,4}\s|```|\||>|[-*]\s|\d+\.\s|---$|\*\*\*$|___$)")


def is_block_start(line):
    """Dong nay mo mot khoi moi (heading/bang/list/code/trich dan) chu khong phai van xuoi."""
    return bool(BLOCK_START.match(line.strip()))


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        add_rich_text(para, text)
        for run in para.runs:
            run.bold = True
        shade(cell._tc.get_or_add_tcPr(), "DCE6F1")
    for row in rows:
        cells = table.add_row().cells
        for i in range(len(header)):
            value = row[i] if i < len(row) else ""
            cells[i].text = ""
            add_rich_text(cells[i].paragraphs[0], value)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PBMS")
    run.font.size = Pt(44)
    run.bold = True
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Parking Building Management System")
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÀI LIỆU CHỨC NĂNG & NGHIỆP VỤ\nBACKEND")
    run.font.size = Pt(22)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Ánh xạ từng nghiệp vụ tới file và dòng code thực thi nó")
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Phạm vi: chỉ Backend (ParkingManagement_BE)\n"
        "Chốt tại nhánh main — 01/08/2026\n"
        "Xuất từ docs/NGHIEP_VU_BACKEND.md trong repository Backend"
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

    h = doc.add_heading("Mục lục", level=1)
    add_toc_field(doc.add_paragraph())
    doc.add_paragraph().runs and None
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def convert(src_path, out_path):
    with open(src_path, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    build_cover(doc)

    i = 0
    skipped_title = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, header, rows)
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level, text = len(m.group(1)), m.group(2)
            if level == 1 and not skipped_title:
                skipped_title = True  # da co trang bia
                i += 1
                continue
            doc.add_heading(text.replace("**", ""), level=max(1, level - 1))
            i += 1
            continue

        # Duong ke ngang
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Trich dan
        if stripped.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            add_rich_text(p, " ".join(x for x in block if x))
            for run in p.runs:
                run.italic = True
            shade(p._p.get_or_add_pPr(), "FFF8E1")
            continue

        # Danh sach -- gop ca cac dong xuong hang cua cung mot muc.
        m = re.match(r"^[-*]\s+(.*)$", stripped) or re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            numbered = bool(re.match(r"^\d+\.\s", stripped))
            text = m.group(2) if numbered else m.group(1)
            i += 1
            cont = [text]
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt or is_block_start(nxt):
                    break
                cont.append(nxt)
                i += 1
            p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
            add_rich_text(p, " ".join(cont))
            continue

        if not stripped:
            i += 1
            continue

        # Doan van thuong. Markdown ngat dong cung trong cung mot doan van chi la
        # de vua be ngang file nguon -- phai GOP lai truoc khi doc dinh dang inline,
        # neu khong mot cap **...** vat qua hai dong se in ra nguyen dau sao.
        block = []
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or is_block_start(nxt):
                break
            if block and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
                break  # dong nay la header cua bang, khong phai van xuoi
            block.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_rich_text(p, " ".join(block))

    doc.save(out_path)
    print("DA XUAT:", out_path)
    print("  headings:", sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading")))
    print("  paragraphs:", len(doc.paragraphs))
    print("  tables:", len(doc.tables))


if __name__ == "__main__":
    convert(SRC, OUT)
